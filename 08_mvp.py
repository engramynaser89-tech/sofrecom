import fitz  # PyMuPDF
import cv2
import os
import re
import requests
import numpy as np
import json
import csv
from docx import Document
from docx.shared import Cm
import sys
import socket
import time
import base64  # for sending images to vision model

# ============================================================
# CONFIG
# ============================================================
pdf_dir = "/Users/ibenferhqt/Downloads/56 Localities (2)"
output_dir = "/Users/ibenferhqt/Downloads/56 Localities (2)/extracted_images10"
legend_dir = "/Users/ibenferhqt/Downloads/56 Localities (2)/cropped_legends10"
report_path = "/Users/ibenferhqt/Downloads/56 Localities (2)/DriveTestReport_10.docx"
log_csv_path = "/Users/ibenferhqt/Downloads/56 Localities (2)/process_log10.csv"

OLLAMA_API = "http://localhost:11434/api/generate"
MODEL_NAME = "llama3.2-vision"  # vision model

# --- OPTIONAL: add RSRP chart (page 6) to report ---
ADD_RSRP_CHART = True  # set to False to disable chart extraction

os.makedirs(output_dir, exist_ok=True)
os.makedirs(legend_dir, exist_ok=True)

# --- Legend cropping parameters (original contour-based logic) ---
min_area = 5000
whiteness_thresh = 236
bottom_right_thresh = 0.5

# --- Page/Image skip lists ---
SKIP_PAGES = [1, 2, 3, 6, 7]     # 1-based page numbers
SKIP_IMAGES = {(5, 1)}           # (page_number, image_number)

# If we want charts, we must NOT skip page 6
if ADD_RSRP_CHART and 6 in SKIP_PAGES:
    SKIP_PAGES.remove(6)

KPI_LIST = ["RSRP", "RSRQ", "SINR"]

# ============================================================
# OLLAMA CONNECTIVITY
# ============================================================
print("🔍 Testing Ollama API connectivity...")

def _check_port(host="localhost", port=11434, timeout=2):
    try:
        with socket.create_connection((host, port), timeout=timeout):
            return True
    except OSError:
        return False

if not _check_port():
    print("❌ Ollama port 11434 is not open on localhost.")
    print("   ➤ Ensure `ollama serve` is running.")
    sys.exit(1)

try:
    print("➡️  Checking available models...")
    tags_resp = requests.get("http://localhost:11434/api/tags", timeout=5)
    if tags_resp.ok:
        tags_data = tags_resp.json()
        model_names = [m.get("name") for m in tags_data.get("models", [])]
        print(f"   ✅ Ollama responded with {len(model_names)} model(s): {model_names}")
    else:
        print(f"   ⚠️ /api/tags returned HTTP {tags_resp.status_code}: {tags_resp.text}")

    print(f"➡️  Sending small generation test to model '{MODEL_NAME}'...")
    test_payload = {"model": MODEL_NAME, "prompt": "ping", "stream": False}
    start_time = time.time()
    test_resp = requests.post(OLLAMA_API, json=test_payload, timeout=20)
    duration = time.time() - start_time

    if test_resp.ok:
        try:
            data = test_resp.json()
            preview = data.get("response", "")[:100]
            print(f"   ✅ Model responded in {duration:.1f}s. Sample output:\n      “{preview}”")
        except Exception as parse_err:
            print(f"   ⚠️ Model responded but JSON parsing failed: {parse_err}\n{test_resp.text}")
    else:
        print(f"❌ Generation request failed: HTTP {test_resp.status_code}")
        print("   Response text:", test_resp.text)
        sys.exit(1)

except requests.exceptions.ConnectTimeout:
    print("❌ Connection timeout: Ollama server did not respond in time.")
    sys.exit(1)
except requests.exceptions.ConnectionError as ce:
    print(f"❌ Connection error: {ce}")
    sys.exit(1)
except requests.exceptions.ReadTimeout:
    print("❌ Read timeout: Ollama took too long to respond.")
    sys.exit(1)
except Exception as e:
    print(f"❌ Unexpected error contacting Ollama: {type(e).__name__}: {e}")
    sys.exit(1)

print("✅ Ollama connectivity test completed successfully.\n")

# ============================================================
# FILENAME → LOCALITY
# ============================================================
def get_locality_from_pdf(filename):
    basename = os.path.basename(filename)

    # STC Data WO pattern, e.g. "20251130_130124_Data_STC WO_2_21042_RG.pdf"
    stc_wo_rg = re.search(r'WO_\d+_(\d+)_RG(?:\.[Pp][Dd][Ff])?$', basename, re.IGNORECASE)
    if stc_wo_rg:
        return stc_wo_rg.group(1)

    # Existing patterns
    end_anchor = re.search(r'_(\d+)_(?:Mobily_)?RS(?:\.[Pp][Dd][Ff])?$', basename, re.IGNORECASE)
    if end_anchor:
        return end_anchor.group(1)

    stc_bi = re.search(r'_(\d+)_BI(?:\.[Pp][Dd][Ff])?$', basename, re.IGNORECASE)
    if stc_bi:
        return stc_bi.group(1)

    stc_plain = re.search(r'_(\d+)(?:\.[Pp][Dd][Ff])$', basename, re.IGNORECASE)
    if stc_plain:
        return stc_plain.group(1)

    mobily_mid = re.search(r'_Mobily_[^_]*_(\d+)_RS', basename, re.IGNORECASE)
    if mobily_mid:
        return mobily_mid.group(1)

    legacy_match = re.search(r'MCIT_july_(\d+)', basename, re.IGNORECASE)
    if legacy_match:
        return legacy_match.group(1)

    return None

# ============================================================
# LEGEND CROPPING
# ============================================================
def crop_legend(image):
    h, w = image.shape[:2]
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    blur = cv2.GaussianBlur(gray, (5, 5), 0)
    _, thresh = cv2.threshold(blur, 249, 255, cv2.THRESH_BINARY)
    contours, _ = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    best_crop, best_score = None, 0
    for cnt in contours:
        x, y, cw, ch = cv2.boundingRect(cnt)
        area = cw * ch
        cx, cy = x + cw / 2, y + ch / 2
        if area < min_area:
            continue
        if cx < w * bottom_right_thresh or cy < h * bottom_right_thresh:
            continue
        crop = image[y:y + ch, x:x + cw]
        mean_val = crop.mean()
        if mean_val > whiteness_thresh:
            score = area * mean_val
            if score > best_score:
                best_crop, best_score = crop, score
    return best_crop

def crop_legend_alt(image):
    h, w = image.shape[:2]
    # 1) Bottom-right ROI
    x1, y1 = int(0.45 * w), int(0.70 * h)
    roi = image[y1:h, x1:w]
    gray = cv2.cvtColor(roi, cv2.COLOR_BGR2GRAY)
    # Emphasize text on light background
    thr = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_MEAN_C,
                                cv2.THRESH_BINARY_INV, 31, 15)
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (7, 7))
    morph = cv2.morphologyEx(thr, cv2.MORPH_CLOSE, kernel, iterations=2)
    contours, _ = cv2.findContours(morph, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

    # pick the largest plausible block
    cand = None
    cand_area = 0
    for c in contours:
        x, y, cw, ch = cv2.boundingRect(c)
        area = cw * ch
        if area > cand_area and cw > 80 and ch > 60:
            cand_area = area
            cand = (x, y, cw, ch)
    if cand is not None:
        x, y, cw, ch = cand
        return roi[y:y+ch, x:x+cw]

    # 2) Fixed slice fallback (bottom-right 35% x 30%)
    x2, y2 = int(0.65 * w), int(0.70 * h)
    fixed = image[y2:h, x2:w]
    return fixed if fixed.size else None

# --- NEW: crop + clean legend (mask out colour boxes) ---
def crop_and_clean_legend(image, debug_path=None):
    """
    Returns a legend crop with coloured swatches masked white,
    but keeps text as intact as possible.
    """
    legend = crop_legend(image)
    if legend is None:
        legend = crop_legend_alt(image)
    if legend is None:
        return None

    orig_legend = legend.copy()

    hsv = cv2.cvtColor(legend, cv2.COLOR_BGR2HSV)

    color_ranges = [
        ((0, 70, 70), (10, 255, 255)),    # Red
        ((11, 70, 70), (25, 255, 255)),   # Orange
        ((26, 70, 70), (35, 255, 255)),   # Yellow
        ((36, 70, 70), (85, 255, 255)),   # Green
        ((86, 70, 70), (130, 255, 255)),  # Blue/Cyan
    ]

    mask = np.zeros(hsv.shape[:2], dtype=np.uint8)
    for lower, upper in color_ranges:
        mask |= cv2.inRange(hsv, np.array(lower), np.array(upper))

    # Less aggressive dilation to avoid erasing text
    kernel = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (5, 5))
    mask = cv2.dilate(mask, kernel, iterations=1)

    cleaned = legend.copy()
    cleaned[mask > 0] = (255, 255, 255)

    cleaned = cv2.bilateralFilter(cleaned, 7, 60, 60)

    h, w = cleaned.shape[:2]
    if max(h, w) < 600:
        cleaned = cv2.resize(cleaned, None, fx=2, fy=2, interpolation=cv2.INTER_CUBIC)

    if debug_path is not None:
        try:
            concat = np.hstack([orig_legend, cleaned])
            cv2.imwrite(debug_path, concat)
        except Exception:
            pass

    return cleaned

# ============================================================
# RSRP CHART CROPPING (PAGE 6)
# ============================================================
def crop_rsrp_chart(image):
    """
    Heuristic crop for the top-left 4G RSRP chart on slide 6.
    Crops slightly less in height (10% off the bottom) to avoid
    including the 4G CQI title.
    """
    h, w = image.shape[:2]

    # Horizontal: same as before (left half-ish)
    x1 = int(0.05 * w)
    x2 = int(0.52 * w)

    # Vertical: top of chart is the same, but reduce bottom from 0.55*h to 0.45*h
    y1 = int(0.12 * h)
    y2 = int(0.45 * h)   # was 0.55 * h → move up by 10% of page height

    # Safety fallback if crop is too small
    if x2 <= x1 + 50 or y2 <= y1 + 50:
        x1, y1 = 0, 0
        x2, y2 = int(0.5 * w), int(0.5 * h)

    return image[y1:y2, x1:x2]

# ============================================================
# KPI BY LAYOUT
# ============================================================
def expected_metric_for(page_num, img_num):
    if page_num == 4 and img_num == 1:
        return "RSRP"
    if page_num == 4 and img_num == 2:
        return "RSRQ"
    if page_num == 5 and img_num == 2:
        return "SINR"
    return None

# ============================================================
# VISION PROMPT BUILDER (SINR BY ROW ORDER)
# ============================================================
def _build_vision_prompt_for_kpi(kpi):
    base_intro = (
        "You are given a cropped legend from an LTE drive-test map. "
        "Each row shows a range label, a sample count in parentheses, "
        "and a percentage value. Ignore the sample counts and read ONLY "
        "the percentage values from the legend in the image."
    )

    if kpi == "RSRP":
        hint = "This is an LTE UE RSRP (dBm) legend."
        key_block = """
Return a JSON object with this exact structure:

{
  "RSRP": {
    "rsrp_lt_-125": <float>,
    "rsrp_-125_to_-115": <float>,
    "rsrp_-115_to_-100": <float>,
    "rsrp_-100_to_-90": <float>,
    "rsrp_-90_to_-80": <float>,
    "rsrp_ge_-80": <float>
  }
}
""".strip()

    elif kpi == "RSRQ":
        hint = "This is an LTE UE RSRQ (dB) legend."
        key_block = """
Return a JSON object with this exact structure:

{
  "RSRQ": {
    "rsrq_lt_-18": <float>,
    "rsrq_-18_to_-14": <float>,
    "rsrq_-14_to_-9": <float>,
    "rsrq_-9_to_-5": <float>,
    "rsrq_ge_-5": <float>
  }
}
""".strip()

    elif kpi == "SINR":
        hint = (
            "This is an LTE UE SINR (dB) legend. There are exactly five rows. "
            "From TOP to BOTTOM, regardless of the exact text or presence of '<', "
            "they always correspond to these SINR ranges:\n"
            "  • Row 1:   SINR < 0 dB\n"
            "  • Row 2:   0 ≤ SINR < 3 dB\n"
            "  • Row 3:   3 ≤ SINR < 13 dB\n"
            "  • Row 4:   13 ≤ SINR < 20 dB\n"
            "  • Row 5:   SINR ≥ 20 dB\n\n"
            "Sometimes the '<' symbol is partially hidden by coloured dots, so "
            "the first row may look like just '0 (460) 72.90 %'. In that case, "
            "treat the FIRST row as 'SINR < 0 dB', the second as '0–3 dB', etc."
        )
        key_block = """
Return a JSON object with this exact structure, mapping each row by its POSITION:

{
  "SINR": {
    "sinr_lt_0":    <float>,  // percentage of the TOP row  (SINR < 0 dB)
    "sinr_0_to_3":  <float>,  // percentage of the second row (0 ≤ SINR < 3 dB)
    "sinr_3_to_13": <float>,  // percentage of the third row  (3 ≤ SINR < 13 dB)
    "sinr_13_to_20":<float>,  // percentage of the fourth row (13 ≤ SINR < 20 dB)
    "sinr_ge_20":   <float>   // percentage of the bottom row (SINR ≥ 20 dB)
  }
}
""".strip()

    else:
        hint = "This is a KPI legend."
        key_block = """
Return a JSON object with this exact structure:

{
  "RSRP": {
    "rsrp_lt_-125": <float>,
    "rsrp_-125_to_-115": <float>,
    "rsrp_-115_to_-100": <float>,
    "rsrp_-100_to_-90": <float>,
    "rsrp_-90_to_-80": <float>,
    "rsrp_ge_-80": <float>
  }
}
""".strip()

    prompt = f"""
{hint}

{base_intro}

{key_block}

Instructions:
- Replace each <float> with the percentage value you READ from the legend in the image.
- For SINR, map the percentages by ROW ORDER exactly as described above
  (do not infer keys from text like "0" or "0 < 3"; use position).
- Do NOT invent values and do NOT reuse example numbers from this prompt.
- The percentages for a given KPI should sum to approximately 100.00.
- Output ONLY the JSON object, with no additional text or markdown.
""".strip()

    return prompt

# ============================================================
# VISION LEGEND PARSER (llama3.2-vision)
# ============================================================
def parse_legends_with_llm(legend_entries):
    """
    legend_entries: list of dicts with keys:
      - metric: "RSRP"/"RSRQ"/"SINR"
      - legend_path: path to cleaned legend image

    Returns:
      (metrics_summary: dict, llm_raw: str)
    """
    if not legend_entries:
        return {}, ""

    metrics_summary = {}
    raw_chunks = []

    for kpi in KPI_LIST:
        entry = next(
            (e for e in legend_entries
             if e.get("metric") == kpi and e.get("legend_path")),
            None
        )
        if not entry:
            continue

        legend_path = entry["legend_path"]

        try:
            with open(legend_path, "rb") as f:
                b64 = base64.b64encode(f.read()).decode("utf-8")
        except Exception as e:
            print(f"   ⚠️ Could not open legend image for {kpi}: {e}")
            continue

        prompt = _build_vision_prompt_for_kpi(kpi)

        payload = {
            "model": MODEL_NAME,
            "prompt": prompt,
            "images": [b64],
            "stream": False,
            "temperature": 0.0,
            "options": {"num_ctx": 4096},
        }

        try:
            r = requests.post(OLLAMA_API, json=payload, timeout=90)
            if not r.ok:
                print(f"❌ Vision API error for {kpi}: {r.status_code} — {r.text}")
                continue

            raw = r.json().get("response", "").strip()
            raw_chunks.append(f"### {kpi}\n{raw}")

            cleaned = raw.strip()
            if cleaned.startswith("```"):
                cleaned = re.sub(r"^```[a-zA-Z]*", "", cleaned, flags=re.MULTILINE).strip("`").strip()

            brace_start = cleaned.find("{")
            brace_end = cleaned.rfind("}")
            if brace_start != -1 and brace_end != -1 and brace_end > brace_start:
                cleaned = cleaned[brace_start:brace_end + 1]

            try:
                parsed = json.loads(cleaned)
            except json.JSONDecodeError:
                print(f"❌ Failed to parse vision JSON for {kpi}. Raw kept in log.")
                continue

            if kpi in parsed and isinstance(parsed[kpi], dict):
                vals = parsed[kpi]
            else:
                vals = parsed

            if isinstance(vals, dict):
                metrics_summary[kpi] = vals
                print(f"   📊 {MODEL_NAME} {kpi} parsed bins: {vals}")
            else:
                print(f"   ⚠️ Parsed output for {kpi} is not a dict, ignored.")

        except Exception as e:
            print(f"❌ Error querying {MODEL_NAME} for {kpi}: {e}")
            continue

    llm_raw = "\n\n".join(raw_chunks)
    return metrics_summary, llm_raw

# ============================================================
# ASSESSMENT LOGIC (RURAL, EXTREME ONLY)
# ============================================================
def _getf(d: dict, k: str) -> float:
    try:
        return float(d.get(k, 0) or 0)
    except Exception:
        return 0.0

def get_rsrp_good_pct(summary):
    rsrp = summary.get("RSRP", {}) or {}
    return (
        _getf(rsrp, "rsrp_-100_to_-90")
        + _getf(rsrp, "rsrp_-90_to_-80")
        + _getf(rsrp, "rsrp_ge_-80")
    )

def get_rsrp_bad_pct(summary):
    rsrp = summary.get("RSRP", {}) or {}
    return (
        _getf(rsrp, "rsrp_lt_-125")
        + _getf(rsrp, "rsrp_-125_to_-115")
        + _getf(rsrp, "rsrp_-115_to_-100")
    )

def determine_assessment(summary):
    """
    Rural logic, focusing only on extreme issues:
      - Coverage rejection: bad RSRP >= 5%
      - Interference: sinr_lt_0 > 30%
      - Quality: rsrq_lt_-18 + rsrq_-18_to_-14 > 30%
    Interference has priority over quality if both are bad.
    """
    rsrp_bad = get_rsrp_bad_pct(summary)

    # 1) Coverage check
    if rsrp_bad >= 5.0:
        return "rejected"

    rsrq = summary.get("RSRQ", {}) or {}
    sinr = summary.get("SINR", {}) or {}

    extreme_rsrq = _getf(rsrq, "rsrq_lt_-18") + _getf(rsrq, "rsrq_-18_to_-14")
    poor_sinr = _getf(sinr, "sinr_lt_0")  # extreme interference only

    # Interference dominates
    if poor_sinr > 30.0:
        return "approved_interference"

    # Quality-only issue
    if extreme_rsrq > 30.0:
        return "approved_quality"

    return "approved"

ASSESSMENT_TEXT = {
    "approved": (
        "Approved\n\n"
        "The coverage footprint meets applicable compliance thresholds, with {rsrp_pct:.2f}% of the measured "
        "area achieving RSRP better than –100 dBm.\n"
        "SINR and RSRQ results indicate generally good service quality, with only limited sections "
        "showing moderate degradation.\n"
        "The locality is approved."
    ),
    "approved_quality": (
        "Approved with comments - quality degradation\n\n"
        "The coverage footprint satisfies applicable compliance thresholds, with {rsrp_pct:.2f}% of the measured "
        "area achieving RSRP better than –100 dBm; however, service quality presents opportunities for enhancement.\n\n"
        "Recommendation:\n"
        "Evaluate site configuration and parameter optimization to improve signal consistency and "
        "reduce potential quality degradation in the service area."
    ),
    "approved_interference": (
        "Approved with comments - interferences\n\n"
        "The coverage footprint satisfies applicable compliance thresholds, with {rsrp_pct:.2f}% of the measured "
        "area achieving RSRP better than –100 dBm. However, SINR and RSRQ measurements indicate widespread "
        "quality degradation, with a high proportion of samples in the low-to-poor range. These conditions "
        "suggest that, despite strong signal strength, user experience may be impacted by inter-site "
        "interference, environmental factors, or site configuration constraints typical of rural deployments.\n\n"
        "Recommendation:\n"
        "Perform targeted optimization of antenna parameters, downtilts, and power settings to improve "
        "signal quality and mitigate interference effects along the covered routes."
    ),
    "rejected": (
        "Rejected\n\n"
        "The coverage footprint does not meet applicable compliance thresholds, with only {rsrp_pct:.2f}% of the "
        "measured area achieving RSRP better than –100 dBm. Several sections exhibit weak or absent signal, "
        "particularly along the southern route segments.\n"
        "The locality is rejected.\n\n"
        "Recommendation:\n"
        "Implement corrective measures, such as additional site deployment, antenna reorientation, or "
        "power adjustments, to extend coverage and address service gaps."
    ),
}

def render_assessment_text(assessment_key, metrics_summary):
    tpl = ASSESSMENT_TEXT.get(assessment_key, "")
    rsrp_pct = get_rsrp_good_pct(metrics_summary)
    if tpl and rsrp_pct is not None:
        return tpl.format(rsrp_pct=rsrp_pct)
    elif tpl:
        return tpl.replace("{rsrp_pct:.2f}%", "a high share")
    else:
        return "No assessment generated."

# --- Consistency check for each KPI ---
def check_kpi_sums(metrics_summary, tolerance=3.0):
    """
    Prints whether each KPI's bins sum to ~100%.
    """
    for kpi in KPI_LIST:
        bins = metrics_summary.get(kpi)
        if not isinstance(bins, dict):
            continue
        total = sum(_getf(bins, key) for key in bins.keys())
        if abs(total - 100.0) <= tolerance:
            print(f"   ✅ {kpi} percentages sum to {total:.2f}% (within ±{tolerance}%).")
        else:
            print(f"   ⚠️ {kpi} percentages sum to {total:.2f}% (outside ±{tolerance}%).")

# ============================================================
# MAIN PROCESS
# ============================================================
locality_data = {}
processed_logs = []

print(f"\n📂 Scanning PDF directory: {pdf_dir}")

for root, _, files in os.walk(pdf_dir):
    for pdf_file in sorted(files):
        if not pdf_file.lower().endswith('.pdf'):
            continue

        log_row = {
            "file_name": pdf_file,
            "locality": "",
            "total_pages": "",
            "images_found": 0,
            "legends_detected": 0,
            "metrics_parsed": 0,
            "assessment_key": "",
            "status": "Failed",
            "reason": "",
            "llm_raw": "",
            "legend_fallback_used": 0,
        }

        try:
            locality = get_locality_from_pdf(pdf_file)
            log_row["locality"] = locality if locality else ""

            if not locality:
                reason = "Filename does not match required locality pattern."
                print(f"⏭️ Skipping {pdf_file}: {reason}")
                log_row["reason"] = reason
                processed_logs.append(log_row)
                continue

            pdf_path = os.path.join(root, pdf_file)
            doc = fitz.open(pdf_path)
            log_row["total_pages"] = len(doc)
            print(f"\n📄 Processing '{pdf_file}' | Locality: '{locality}' | Total pages: {len(doc)}")

            images_info = []
            legend_entries = []
            rsrp_chart_path = None  # NEW: per-locality RSRP chart

            for page_index in range(1, len(doc)):
                page_num = page_index + 1

                # Normal skip logic
                if page_num in SKIP_PAGES:
                    print(f"   ⏭️  Skipping page {page_num} (skip list)")
                    continue

                page = doc.load_page(page_index)

                # --- NEW: handle page 6 via full-page render for RSRP chart ---
                if ADD_RSRP_CHART and page_num == 6 and rsrp_chart_path is None:
                    try:
                        # render at higher resolution
                        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                        img_mode = "RGBA" if pix.alpha else "RGB"
                        img_np = np.frombuffer(pix.samples, dtype=np.uint8)
                        img_np = img_np.reshape(pix.height, pix.width, pix.n)

                        if pix.n == 4:  # RGBA -> BGR
                            page_img = cv2.cvtColor(img_np, cv2.COLOR_RGBA2BGR)
                        else:  # RGB -> BGR
                            page_img = cv2.cvtColor(img_np, cv2.COLOR_RGB2BGR)

                        chart_img = crop_rsrp_chart(page_img)
                        if chart_img is not None and chart_img.size > 0:
                            chart_filename = f"chart_RSRP_{locality}_p{page_num}.png"
                            chart_path = os.path.join(output_dir, chart_filename)
                            cv2.imwrite(chart_path, chart_img)
                            rsrp_chart_path = chart_path
                            print(f"   📈 RSRP chart extracted from page {page_num} → {chart_path}")
                        else:
                            print(f"   ⚠️ RSRP chart crop empty on page {page_num}.")
                    except Exception as e:
                        print(f"   ⚠️ Failed to render/crop RSRP chart on page {page_num}: {e}")

                    # We don't want to treat page 6 as a legend page
                    print("   ℹ️ Page 6 processed for charts; skipping legend extraction.")
                    continue

                # --- For other pages, process embedded images as before ---
                images = page.get_images(full=True)
                print(f"   🧩 Page {page_num}: {len(images)} images found")
                log_row["images_found"] += len(images)

                for img_index, img in enumerate(images):
                    img_num = img_index + 1
                    if (page_num, img_num) in SKIP_IMAGES:
                        print(f"      ⏭️  Skipping image {img_num} on page {page_num} (skip list)")
                        continue

                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]
                    image_filename = f"{locality}_page{page_num}_img{img_num}_{os.path.splitext(pdf_file)[0]}.{image_ext}"
                    image_path = os.path.join(output_dir, image_filename)
                    with open(image_path, "wb") as f:
                        f.write(image_bytes)
                    print(f"      💾 Saved image: {image_path}")

                    image_cv = cv2.imdecode(np.frombuffer(image_bytes, np.uint8), cv2.IMREAD_COLOR)
                    if image_cv is None:
                        print(f"      ❌ Could not decode {image_path}, removing.")
                        os.remove(image_path)
                        continue

                    # Legend pipeline (maps on pages 4 & 5)
                    debug_legend_path = os.path.join(
                        legend_dir,
                        f"debug_{locality}_p{page_num}_i{img_num}.jpg"
                    )

                    legend_img = crop_and_clean_legend(image_cv, debug_path=debug_legend_path)
                    if legend_img is None:
                        print("      ❌ No legend found after cleaning, removing image.")
                        os.remove(image_path)
                        continue

                    log_row["legends_detected"] += 1

                    legend_filename = f"legend_{os.path.basename(image_path)}"
                    legend_path = os.path.join(legend_dir, legend_filename)
                    try:
                        cv2.imwrite(legend_path, legend_img)
                        print(f"      ✂️  Cleaned legend saved → {legend_path}")
                    except Exception as e:
                        print(f"      ❌ Failed to save legend crop: {e}")
                        continue

                    assigned_metric = expected_metric_for(page_num, img_num)
                    if assigned_metric is None:
                        print(f"      ℹ️  No KPI mapping for page {page_num}, img {img_num} → legend kept but not sent to vision model.")
                    else:
                        legend_entries.append({
                            'metric': assigned_metric,
                            'legend_path': legend_path,
                        })

                    images_info.append({
                        'img_path': image_path,
                        'metric_assigned': assigned_metric,
                        'legend_path': legend_path,
                        'fallback_used': False,
                        'page_num': page_num,
                        'img_num': img_num,
                    })

            doc.close()

            if images_info:
                print(f"🧠 Sending legends to vision model '{MODEL_NAME}' for parsing (by KPI)...")
                metrics_summary, llm_raw = parse_legends_with_llm(legend_entries)
                log_row["llm_raw"] = llm_raw or ""

                if isinstance(metrics_summary, dict) and metrics_summary:
                    log_row["metrics_parsed"] = len(metrics_summary.keys())

                # Consistency check on sums
                if metrics_summary:
                    check_kpi_sums(metrics_summary)

                if metrics_summary:
                    assessment_key = determine_assessment(metrics_summary)
                else:
                    assessment_key = "rejected"
                log_row["assessment_key"] = assessment_key

                locality_data[locality] = {
                    'images_info': images_info,
                    'metrics': metrics_summary,
                    'assessment_key': assessment_key,
                    'rsrp_chart_path': rsrp_chart_path,
                }

                log_row["status"] = "Succeeded" if metrics_summary else "Failed"
                if not metrics_summary:
                    log_row["reason"] = "No metrics parsed from vision model."
                print(f"✅ Locality {locality} assessment → {assessment_key.upper()}")
            else:
                reason = f"No relevant images/legends extracted."
                print(f"⚠️ {reason} From {pdf_file} ({locality}) - Not added to report.")
                log_row["reason"] = reason

        except Exception as e:
            log_row["reason"] = f"Exception: {e}"
            print(f"❌ Error processing {pdf_file}: {e}")

        processed_logs.append(log_row)

print("\n✅ All eligible PDFs processed.\n")

# ============================================================
# BUILD WORD REPORT
# ============================================================
document = Document()
if locality_data:
    for locality, data in locality_data.items():
        document.add_heading(f"Locality: {locality}", level=1)

        # pick first occurrence per KPI using assigned metric
        first_per_kpi = {}
        for item in data['images_info']:
            k = item.get('metric_assigned') or 'Metric'
            if k in KPI_LIST and k not in first_per_kpi:
                first_per_kpi[k] = item

        selected = [first_per_kpi[k] for k in KPI_LIST if k in first_per_kpi]
        if selected:
            table = document.add_table(rows=2, cols=len(selected))
            hdr_cells = table.rows[0].cells
            img_cells = table.rows[1].cells
            for idx, imginfo in enumerate(selected):
                hdr_cells[idx].text = imginfo.get('metric_assigned') or 'Metric'
                try:
                    p = img_cells[idx].paragraphs[0]
                    run = p.add_run()
                    run.add_picture(imginfo['img_path'], width=Cm(5))
                except Exception as e:
                    print(f"    ❌ Failed to add picture: {e}")

        metrics_summary = data.get('metrics', {})
        assessment_key = data.get('assessment_key', 'approved')
        assessment_text = render_assessment_text(assessment_key, metrics_summary)

        document.add_heading("Coverage Assessment", level=2)
        for block in assessment_text.split("\n\n"):
            document.add_paragraph(block)

        # --- OPTIONAL: insert RSRP chart just after assessment ---
        if ADD_RSRP_CHART and data.get("rsrp_chart_path"):
            try:
                document.add_heading("4G RSRP KPI Chart", level=2)
                p = document.add_paragraph()
                run = p.add_run()
                run.add_picture(data["rsrp_chart_path"], width=Cm(12))
            except Exception as e:
                print(f"    ⚠️ Failed to add RSRP chart for locality {locality}: {e}")

        document.add_page_break()
else:
    document.add_heading("Drive Test Report", level=1)
    document.add_paragraph("No localities were added to the report due to missing legends/metrics or processing errors.")

document.save(report_path)
print(f"\n✅ Final report generated at: {report_path}\n")

# ============================================================
# CSV LOG
# ============================================================
log_fields = [
    "file_name",
    "locality",
    "total_pages",
    "images_found",
    "legends_detected",
    "metrics_parsed",
    "assessment_key",
    "status",
    "reason",
    "llm_raw",
    "legend_fallback_used",
]
try:
    with open(log_csv_path, mode="w", newline="", encoding="utf-8") as csvfile:
        writer = csv.DictWriter(csvfile, fieldnames=log_fields)
        writer.writeheader()
        for row in processed_logs:
            writer.writerow(row)
    print(f"🧾 Process log saved to: {log_csv_path}")
except Exception as e:
    print(f"❌ Failed to write process log CSV: {e}")