import fitz  # PyMuPDF
import cv2
import pytesseract
import os
import re
import requests
import numpy as np
import json
import csv
from docx import Document
from docx.shared import Cm
import sys
import socket
import time

# --- Configuration ---
pdf_dir = "/Users/ibenferhqt/Downloads/56 Localities (1)"
output_dir = "/Users/ibenferhqt/Downloads/56 Localities (1)/extracted_images"
legend_dir = "/Users/ibenferhqt/Downloads/56 Localities (1)/cropped_legends"
report_path = "/Users/ibenferhqt/Downloads/56 Localities (1)/DriveTestReport_2.docx"
log_csv_path = "/Users/ibenferhqt/Downloads/56 Localities (1)/process_log.csv"

OLLAMA_API = "http://localhost:11434/api/generate"
MODEL_NAME = "gpt-oss:20b"

os.makedirs(output_dir, exist_ok=True)
os.makedirs(legend_dir, exist_ok=True)

# --- Legend cropping parameters (original contour-based logic) ---
min_area = 5000
whiteness_thresh = 236
bottom_right_thresh = 0.5

# --- Page/Image skip lists ---
SKIP_PAGES = [1, 2, 3, 6, 7]     # 1-based page numbers (your latest setting)
SKIP_IMAGES = {(5, 1)}           # (page_number, image_number)

KPI_LIST = ["RSRP", "RSRQ", "SINR"]

# --- Ollama connectivity diagnostics (verbose) ---
print("🔍 Testing Ollama API connectivity...")

def _check_port(host="localhost", port=11434, timeout=2):
    try:
        with socket.create_connection((host, port), timeout=timeout):
            return True
    except OSError:
        return False

if not _check_port():
    print("❌ Ollama port 11434 is not open on localhost.")
    print("   ➤ Ensure `ollama serve` is running.")
    sys.exit(1)

try:
    print("➡️  Checking available models...")
    tags_resp = requests.get("http://localhost:11434/api/tags", timeout=5)
    if tags_resp.ok:
        tags_data = tags_resp.json()
        model_names = [m.get("name") for m in tags_data.get("models", [])]
        print(f"   ✅ Ollama responded with {len(model_names)} model(s): {model_names}")
    else:
        print(f"   ⚠️ /api/tags returned HTTP {tags_resp.status_code}: {tags_resp.text}")

    print(f"➡️  Sending small generation test to model '{MODEL_NAME}'...")
    test_payload = {"model": MODEL_NAME, "prompt": "ping", "stream": False}
    start_time = time.time()
    test_resp = requests.post(OLLAMA_API, json=test_payload, timeout=20)
    duration = time.time() - start_time

    if test_resp.ok:
        try:
            data = test_resp.json()
            preview = data.get("response", "")[:100]
            print(f"   ✅ Model responded in {duration:.1f}s. Sample output:\n      “{preview}”")
        except Exception as parse_err:
            print(f"   ⚠️ Model responded but JSON parsing failed: {parse_err}\n{test_resp.text}")
    else:
        print(f"❌ Generation request failed: HTTP {test_resp.status_code}")
        print("   Response text:", test_resp.text)
        sys.exit(1)

except requests.exceptions.ConnectTimeout:
    print("❌ Connection timeout: Ollama server did not respond in time.")
    sys.exit(1)
except requests.exceptions.ConnectionError as ce:
    print(f"❌ Connection error: {ce}")
    sys.exit(1)
except requests.exceptions.ReadTimeout:
    print("❌ Read timeout: Ollama took too long to respond.")
    sys.exit(1)
except Exception as e:
    print(f"❌ Unexpected error contacting Ollama: {type(e).__name__}: {e}")
    sys.exit(1)

print("✅ Ollama connectivity test completed successfully.\n")

# --- Filename parsing (unchanged) ---
import os
import re

import os
import re

def get_locality_from_pdf(filename):
    basename = os.path.basename(filename)

    # NEW: STC Data WO pattern, e.g. "20251130_130124_Data_STC WO_2_21042_RG.pdf"
    # This looks for "WO_<some number>_<locality>_RG.pdf" anywhere near the end.
    stc_wo_rg = re.search(r'WO_\d+_(\d+)_RG(?:\.[Pp][Dd][Ff])?$', basename, re.IGNORECASE)
    if stc_wo_rg:
        return stc_wo_rg.group(1)

    # Existing patterns
    end_anchor = re.search(r'_(\d+)_(?:Mobily_)?RS(?:\.[Pp][Dd][Ff])?$', basename, re.IGNORECASE)
    if end_anchor:
        return end_anchor.group(1)

    stc_bi = re.search(r'_(\d+)_BI(?:\.[Pp][Dd][Ff])?$', basename, re.IGNORECASE)
    if stc_bi:
        return stc_bi.group(1)

    stc_plain = re.search(r'_(\d+)(?:\.[Pp][Dd][Ff])$', basename, re.IGNORECASE)
    if stc_plain:
        return stc_plain.group(1)

    mobily_mid = re.search(r'_Mobily_[^_]*_(\d+)_RS', basename, re.IGNORECASE)
    if mobily_mid:
        return mobily_mid.group(1)

    legacy_match = re.search(r'MCIT_july_(\d+)', basename, re.IGNORECASE)
    if legacy_match:
        return legacy_match.group(1)

    return None

# --- Original legend cropper (contour-based) ---
def crop_legend(image):
    h, w = image.shape[:2]
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    blur = cv2.GaussianBlur(gray, (5, 5), 0)
    _, thresh = cv2.threshold(blur, 249, 255, cv2.THRESH_BINARY)
    contours, _ = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    best_crop, best_score = None, 0
    for cnt in contours:
        x, y, cw, ch = cv2.boundingRect(cnt)
        area = cw * ch
        cx, cy = x + cw / 2, y + ch / 2
        if area < min_area:
            continue
        if cx < w * bottom_right_thresh or cy < h * bottom_right_thresh:
            continue
        crop = image[y:y + ch, x:x + cw]
        mean_val = crop.mean()
        if mean_val > whiteness_thresh:
            score = area * mean_val
            if score > best_score:
                best_crop, best_score = crop, score
    return best_crop

# --- Alternate legend cropper (ROI + morphology, then fixed slice fallback) ---
def crop_legend_alt(image):
    h, w = image.shape[:2]
    # 1) Bottom-right ROI
    x1, y1 = int(0.45 * w), int(0.70 * h)
    roi = image[y1:h, x1:w]
    gray = cv2.cvtColor(roi, cv2.COLOR_BGR2GRAY)
    # Emphasize text on light background
    thr = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_MEAN_C,
                                cv2.THRESH_BINARY_INV, 31, 15)
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (7, 7))
    morph = cv2.morphologyEx(thr, cv2.MORPH_CLOSE, kernel, iterations=2)
    contours, _ = cv2.findContours(morph, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

    # pick the largest plausible block
    cand = None
    cand_area = 0
    for c in contours:
        x, y, cw, ch = cv2.boundingRect(c)
        area = cw * ch
        if area > cand_area and cw > 80 and ch > 60:
            cand_area = area
            cand = (x, y, cw, ch)
    if cand is not None:
        x, y, cw, ch = cand
        return roi[y:y+ch, x:x+cw]

    # 2) Fixed slice fallback (bottom-right 35% x 30%)
    x2, y2 = int(0.65 * w), int(0.70 * h)
    fixed = image[y2:h, x2:w]
    return fixed if fixed.size else None

# --- OCR (unchanged) ---
def run_ocr_strong(image):
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    tries = [
        dict(thresh=165, psm=6),
        dict(thresh=180, psm=6),
        dict(thresh=150, psm=6),
        dict(thresh=165, psm=7),
        dict(thresh=165, psm=11),
        dict(thresh=200, psm=6),
    ]
    for i, setting in enumerate(tries):
        _, binarized = cv2.threshold(gray, setting['thresh'], 255, cv2.THRESH_BINARY)
        config = f"--psm {setting['psm']}"
        text = pytesseract.image_to_string(binarized, config=config).strip()
        if re.search(r"RSRP", text, re.I):
            print(f"      ✔️ RSRP legend found with OCR setting {i+1}")
            return text
    print("      ⚠️ RSRP not found with all settings, returning last attempt")
    _, binarized = cv2.threshold(gray, 165, 255, cv2.THRESH_BINARY)
    return pytesseract.image_to_string(binarized, config="--psm 6").strip()

def run_ocr_simple(image):
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    _, binarized = cv2.threshold(gray, 165, 255, cv2.THRESH_BINARY)
    return pytesseract.image_to_string(binarized, config="--psm 6").strip()

def get_metric_from_text(text):
    for kpi in KPI_LIST:
        if kpi in text.upper():
            return kpi
    for kpi in KPI_LIST:
        if kpi.lower() in text.lower():
            return kpi
    return "Metric"

# --- Fixed KPI by layout ---
def expected_metric_for(page_num, img_num):
    if page_num == 4 and img_num == 1:
        return "RSRP"
    if page_num == 4 and img_num == 2:
        return "RSRQ"
    if page_num == 5 and img_num == 2:
        return "SINR"
    return None  # other images (if any) can still rely on OCR guess

# --- LLM parser: NEW legend schema ---
def parse_legends_with_llm(legend_entries):
    if not legend_entries:
        return {}, ""
    legend_blocks = []
    for e in legend_entries:
        legend_blocks.append(f"KPI: {e['metric']}\nOCR:\n<<<\n{e['ocr_text']}\n>>>")

    prompt = """
You are given OCR text from map legends for KPIs: RSRP, RSRQ, and SINR.
Convert each KPI's legend into a SINGLE JSON object with these exact keys (omit KPIs not present).

RSRP (dBm):
  rsrp_lt_-125
  rsrp_-125_to_-115
  rsrp_-115_to_-100
  rsrp_-100_to_-90
  rsrp_-90_to_-80
  rsrp_ge_-80

RSRQ (dB):
  rsrq_lt_-18
  rsrq_-18_to_-14
  rsrq_-14_to_-9
  rsrq_-9_to_-5
  rsrq_ge_-5

SINR (dB):
  sinr_lt_0
  sinr_0_to_3
  sinr_3_to_13
  sinr_13_to_20
  sinr_ge_20

Conventions:
- Values are percentages (float, no % symbol).
- Each KPI totals ≈ 100%.
- Return only the JSON object, no commentary or code fences.
""".strip()

    payload = {
        "model": MODEL_NAME,
        "prompt": f"{prompt}\n\nLegend data:\n" + "\n\n".join(legend_blocks),
        "stream": False,
    }

    try:
        r = requests.post(OLLAMA_API, json=payload, timeout=120)
        if not r.ok:
            print(f"❌ Ollama API error: {r.status_code} — {r.text}")
            return {}, ""
        raw = r.json().get("response", "").strip()
        cleaned = raw.strip()
        if cleaned.startswith("```"):
            cleaned = re.sub(r"^```[a-zA-Z]*", "", cleaned).strip("`")
        try:
            return json.loads(cleaned), cleaned
        except json.JSONDecodeError:
            print("❌ Failed to parse LLM JSON. Raw below:")
            print(cleaned)
            return {}, cleaned
    except Exception as e:
        print(f"❌ Error querying LLM: {e}")
        return {}, ""

# --- Acceptance thresholds (dual-condition) ---
def _getf(d: dict, k: str) -> float:
    try:
        return float(d.get(k, 0) or 0)
    except Exception:
        return 0.0

def get_rsrp_good_pct(summary):
    rsrp = summary.get("RSRP", {}) or {}
    return _getf(rsrp, "rsrp_-100_to_-90") + _getf(rsrp, "rsrp_-90_to_-80") + _getf(rsrp, "rsrp_ge_-80")

def get_rsrp_bad_pct(summary):
    rsrp = summary.get("RSRP", {}) or {}
    return _getf(rsrp, "rsrp_lt_-125") + _getf(rsrp, "rsrp_-125_to_-115") + _getf(rsrp, "rsrp_-115_to_-100")

def determine_assessment(summary):
    rsrp_good = get_rsrp_good_pct(summary)
    rsrp_bad = get_rsrp_bad_pct(summary)
    if not (rsrp_good >= 95.0 or rsrp_bad < 5.0):
        return "rejected"
    rsrq = summary.get("RSRQ", {}) or {}
    sinr = summary.get("SINR", {}) or {}
    low_quality = _getf(rsrq, "rsrq_-14_to_-9") + _getf(rsrq, "rsrq_-18_to_-14")
    poor_sinr = _getf(sinr, "sinr_lt_0") + _getf(sinr, "sinr_0_to_3")
    if poor_sinr >= 20 or low_quality >= 20:
        return "approved_interference"
    if poor_sinr >= 10 or low_quality >= 10:
        return "approved_quality"
    return "approved"

# --- Assessment text (unchanged) ---
ASSESSMENT_TEXT = {
    "approved": (
        "Approved\n\n"
        "The coverage footprint meets applicable compliance thresholds, with {rsrp_pct:.2f}% of the measured "
        "area achieving RSRP better than –100 dBm.\n"
        "SINR and RSRQ results indicate generally good service quality, with only limited sections "
        "showing moderate degradation.\n"
        "The locality is approved."
    ),
    "approved_quality": (
        "Approved with comments - quality degradation\n\n"
        "The coverage footprint satisfies applicable compliance thresholds, with {rsrp_pct:.2f}% of the measured "
        "area achieving RSRP better than –100 dBm; however, service quality presents opportunities for enhancement.\n\n"
        "Recommendation:\n"
        "Evaluate site configuration and parameter optimization to improve signal consistency and "
        "reduce potential quality degradation in the service area."
    ),
    "approved_interference": (
        "Approved with comments - interferences\n\n"
        "The coverage footprint satisfies applicable compliance thresholds, with {rsrp_pct:.2f}% of the measured "
        "area achieving RSRP better than –100 dBm. However, SINR and RSRQ measurements indicate widespread "
        "quality degradation, with a high proportion of samples in the low-to-poor range. These conditions "
        "suggest that, despite strong signal strength, user experience may be impacted by inter-site "
        "interference, environmental factors, or site configuration constraints typical of rural deployments.\n\n"
        "Recommendation:\n"
        "Perform targeted optimization of antenna parameters, downtilts, and power settings to improve "
        "signal quality and mitigate interference effects along the covered routes."
    ),
    "rejected": (
        "Rejected\n\n"
        "The coverage footprint does not meet applicable compliance thresholds, with only {rsrp_pct:.2f}% of the "
        "measured area achieving RSRP better than –100 dBm. Several sections exhibit weak or absent signal, "
        "particularly along the southern route segments.\n"
        "The locality is rejected.\n\n"
        "Recommendation:\n"
        "Implement corrective measures, such as additional site deployment, antenna reorientation, or "
        "power adjustments, to extend coverage and address service gaps."
    ),
}

def render_assessment_text(assessment_key, metrics_summary):
    tpl = ASSESSMENT_TEXT.get(assessment_key, "")
    rsrp_pct = get_rsrp_good_pct(metrics_summary)
    if tpl and rsrp_pct is not None:
        return tpl.format(rsrp_pct=rsrp_pct)
    elif tpl:
        return tpl.replace("{rsrp_pct:.2f}%", "a high share")
    else:
        return "No assessment generated."

# --- Main Process ---
locality_data = {}
processed_logs = []

print(f"\n📂 Scanning PDF directory: {pdf_dir}")

for root, _, files in os.walk(pdf_dir):
    for pdf_file in sorted(files):
        if not pdf_file.lower().endswith('.pdf'):
            continue

        log_row = {
            "file_name": pdf_file,
            "locality": "",
            "total_pages": "",
            "images_found": 0,
            "legends_detected": 0,
            "metrics_parsed": 0,
            "assessment_key": "",
            "status": "Failed",
            "reason": "",
            "llm_raw": "",
            "ocr_mismatch_count": 0,
            "legend_fallback_used": 0,
        }

        try:
            locality = get_locality_from_pdf(pdf_file)
            log_row["locality"] = locality if locality else ""

            if not locality:
                reason = "Filename does not match required locality pattern."
                print(f"⏭️ Skipping {pdf_file}: {reason}")
                log_row["reason"] = reason
                processed_logs.append(log_row)
                continue

            pdf_path = os.path.join(root, pdf_file)
            doc = fitz.open(pdf_path)
            log_row["total_pages"] = len(doc)
            print(f"\n📄 Processing '{pdf_file}' | Locality: '{locality}' | Total pages: {len(doc)}")

            images_info = []
            legend_entries = []

            for page_index in range(1, len(doc)):  # your original loop starts at second page
                page_num = page_index + 1  # 1-based

                if page_num in SKIP_PAGES:
                    print(f"   ⏭️  Skipping page {page_num} (skip list)")
                    continue

                page = doc.load_page(page_index)
                images = page.get_images(full=True)
                print(f"   🧩 Page {page_num}: {len(images)} images found")
                log_row["images_found"] += len(images)

                for img_index, img in enumerate(images):
                    img_num = img_index + 1
                    if (page_num, img_num) in SKIP_IMAGES:
                        print(f"      ⏭️  Skipping image {img_num} on page {page_num} (skip list)")
                        continue

                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]
                    image_filename = f"{locality}_page{page_num}_img{img_num}_{os.path.splitext(pdf_file)[0]}.{image_ext}"
                    image_path = os.path.join(output_dir, image_filename)
                    with open(image_path, "wb") as f:
                        f.write(image_bytes)
                    print(f"      💾 Saved image: {image_path}")

                    image_cv = cv2.imdecode(np.frombuffer(image_bytes, np.uint8), cv2.IMREAD_COLOR)
                    if image_cv is None:
                        print(f"      ❌ Could not decode {image_path}, removing.")
                        os.remove(image_path)
                        continue

                    # Primary legend detection
                    legend_img = crop_legend(image_cv)
                    fallback_used = False

                    # Fallback legend detection if needed
                    if legend_img is None:
                        print(f"      ⚠️ Primary legend detector failed → trying alternate method...")
                        legend_img = crop_legend_alt(image_cv)
                        if legend_img is not None:
                            fallback_used = True
                            log_row["legend_fallback_used"] += 1
                            print("      ✅ Alternate legend crop succeeded.")
                        else:
                            print(f"      ❌ Legend not detected after alternate method, removing image.")
                            os.remove(image_path)
                            continue

                    # Count a detected legend only when crop exists
                    log_row["legends_detected"] += 1

                    # OCR
                    if page_num == 4 and img_num == 1:
                        ocr_text = run_ocr_strong(legend_img)
                    else:
                        ocr_text = run_ocr_simple(legend_img)

                    # KPI assignment by layout (no guessing)
                    assigned_metric = expected_metric_for(page_num, img_num) or get_metric_from_text(ocr_text)
                    ocr_metric = get_metric_from_text(ocr_text)
                    ocr_match = (assigned_metric == ocr_metric and assigned_metric in KPI_LIST)

                    if not ocr_match and assigned_metric in KPI_LIST:
                        log_row["ocr_mismatch_count"] += 1
                        print(f"      ⚠️ OCR KPI '{ocr_metric}' ≠ expected '{assigned_metric}' — will assign by layout anyway.")

                    legend_filename = f"legend_{os.path.basename(image_path)}"
                    legend_path = os.path.join(legend_dir, legend_filename)
                    try:
                        cv2.imwrite(legend_path, legend_img)
                        print(f"      ✂️  Legend cropped → {legend_path} ({'fallback' if fallback_used else 'primary'})")
                    except Exception as e:
                        print(f"      ❌ Failed to save legend crop: {e}")

                    # For LLM, we pass the assigned metric (layout) but we keep the OCR metric for traceability
                    images_info.append({
                        'img_path': image_path,
                        'metric_assigned': assigned_metric,
                        'metric_ocr': ocr_metric,
                        'ocr_match': ocr_match,
                        'ocr_text': ocr_text,
                        'legend_path': legend_path,
                        'fallback_used': fallback_used,
                        'page_num': page_num,
                        'img_num': img_num,
                    })
                    legend_entries.append({'metric': assigned_metric, 'ocr_text': ocr_text})
                    print(f"      🔎 OCR KPI={ocr_metric} | assigned={assigned_metric} | match={ocr_match} | text: {ocr_text[:80].replace('\\n',' ')}")

            doc.close()

            if images_info:
                print(f"🧠 Sending {len(legend_entries)} legend(s) to LLM for parsing...")
                metrics_summary, llm_raw = parse_legends_with_llm(legend_entries)
                log_row["llm_raw"] = llm_raw or ""

                if isinstance(metrics_summary, dict) and metrics_summary:
                    log_row["metrics_parsed"] = len(metrics_summary.keys())

                assessment_key = determine_assessment(metrics_summary)
                log_row["assessment_key"] = assessment_key

                locality_data[locality] = {
                    'images_info': images_info,
                    'metrics': metrics_summary,
                    'assessment_key': assessment_key,
                }

                log_row["status"] = "Succeeded"
                log_row["reason"] = ""
                print(f"✅ Locality {locality} assessment → {assessment_key.upper()}")
            else:
                reason = f"No relevant images/legends extracted."
                print(f"⚠️ {reason} From {pdf_file} ({locality}) - Not added to report.")
                log_row["reason"] = reason

        except Exception as e:
            log_row["reason"] = f"Exception: {e}"
            print(f"❌ Error processing {pdf_file}: {e}")

        processed_logs.append(log_row)

print("\n✅ All eligible PDFs processed.\n")

# --- Step 2: Build report (exactly one image per KPI if found) ---
document = Document()
if locality_data:
    for locality, data in locality_data.items():
        document.add_heading(f"Locality: {locality}", level=1)

        # pick first occurrence per KPI using assigned metric
        first_per_kpi = {}
        for item in data['images_info']:
            k = item.get('metric_assigned') or item.get('metric_ocr') or 'Metric'
            if k in KPI_LIST and k not in first_per_kpi:
                first_per_kpi[k] = item

        selected = [first_per_kpi[k] for k in KPI_LIST if k in first_per_kpi]
        if selected:
            table = document.add_table(rows=2, cols=len(selected))
            hdr_cells = table.rows[0].cells
            img_cells = table.rows[1].cells
            for idx, imginfo in enumerate(selected):
                hdr_cells[idx].text = imginfo.get('metric_assigned') or imginfo.get('metric_ocr') or 'Metric'
                try:
                    p = img_cells[idx].paragraphs[0]
                    run = p.add_run()
                    run.add_picture(imginfo['img_path'], width=Cm(5))
                except Exception as e:
                    print(f"    ❌ Failed to add picture: {e}")

        metrics_summary = data.get('metrics', {})
        assessment_key = data.get('assessment_key', 'approved')
        assessment_text = render_assessment_text(assessment_key, metrics_summary)

        document.add_heading("Coverage Assessment", level=2)
        for block in assessment_text.split("\n\n"):
            document.add_paragraph(block)

        document.add_page_break()
else:
    document.add_heading("Drive Test Report", level=1)
    document.add_paragraph("No localities were added to the report due to missing legends/metrics or processing errors.")

document.save(report_path)
print(f"\n✅ Final report generated at: {report_path}\n")

# --- Step 3: Write CSV log of processing results ---
log_fields = [
    "file_name",
    "locality",
    "total_pages",
    "images_found",
    "legends_detected",
    "metrics_parsed",
    "assessment_key",
    "status",
    "reason",
    "llm_raw",
    "ocr_mismatch_count",
    "legend_fallback_used",
]
try:
    with open(log_csv_path, mode="w", newline="", encoding="utf-8") as csvfile:
        writer = csv.DictWriter(csvfile, fieldnames=log_fields)
        writer.writeheader()
        for row in processed_logs:
            writer.writerow(row)
    print(f"🧾 Process log saved to: {log_csv_path}")
except Exception as e:
    print(f"❌ Failed to write process log CSV: {e}")